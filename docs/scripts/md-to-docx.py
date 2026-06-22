"""
Markdown → DOCX 转换器
- 替换 mermaid 块 → 渲染好的 PNG 引用
- 替换 ASCII 框图 → 真实页面截图(智能匹配)
- 支持:标题、段落、列表、表格、代码块、引用、链接、加粗/斜体
- 输出: docs/sop/<name>.docx
"""
import re
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SOP_DIR = Path("docs/sop")
IMG_MERMAID = SOP_DIR / "_images" / "mermaid"
IMG_SCREENSHOT = SOP_DIR / "_images" / "screenshots"

# 路由 → 截图文件 key(包含更多关键词,提高匹配)
ROUTE_PNG = {
    # 高频关键词(短一些也能匹配)
    "dashboard": "01-dashboard",
    "仪表盘": "01-dashboard",
    "首页": "01-dashboard",
    "款式管理": "03-styles",
    "款式": "03-styles",
    "styles": "03-styles",
    "订单": "03-styles",
    "分色分尺码": "04-style-color-size",
    "分色": "04-style-color-size",
    "color-size": "04-style-color-size",
    "面料装载清单": "05-fabric-list",
    "面料装柜": "05-fabric-list",
    "面料": "05-fabric-list",
    "fabric": "05-fabric-list",
    "缝制车间管理": "06-sewing-workshop",
    "缝制车间": "06-sewing-workshop",
    "workshop": "06-sewing-workshop",
    "三层树": "06-sewing-workshop",
    "前置车间": "07-pre-workshop-output",
    "基础数据": "02-basic-data",
    "basic-data": "02-basic-data",
    "总计划": "09-main-plan",
    "主计划": "09-main-plan",
    "main-plan": "09-main-plan",
    "计划管理首页": "08-plan-management",
    "裁剪排程": "10-cutting-schedule",
    "裁剪": "10-cutting-schedule",
    "cutting": "10-cutting-schedule",
    "二次加工首页": "11-secondary-home",
    "二次": "11-secondary-home",
    "secondary": "11-secondary-home",
    "印花计划": "12-printing-plan",
    "刺绣计划": "13-embroidery-plan",
    "模板计划": "14-template-plan",
    "烫标计划": "15-ironing-plan",
    "缝制首页": "16-sewing-home",
    "缝制排程详情": "17-sewing-plan-detail",
    "目视化排程": "18-sewing-visual",
    "排程总览": "18-sewing-visual",
    "视觉": "18-sewing-visual",
    "报工首页": "19-dispatch-home",
    "报工汇总": "20-dispatch-report",
    "报工总览": "19-dispatch-home",
    "裁剪一检": "21-cutting-dispatch",
    "裁剪二检": "22-cutting-second-dispatch",
    "裁剪报工": "21-cutting-dispatch",
    "印花报工": "23-printing-dispatch",
    "刺绣报工": "24-embroidery-dispatch",
    "模板报工": "25-template-dispatch",
    "烫标报工": "26-ironing-dispatch",
    "缝制报工": "27-sewing-dispatch-select",
    "报工详情": "28-sewing-dispatch-detail",
    "本车间": "28-sewing-dispatch-detail",
    "报工员": "19-dispatch-home",
    "报工生成的入库": "30-warehouse-cutting-piece",
    "面料入库": "30-warehouse-cutting-piece",
    "裁片入库": "30-warehouse-cutting-piece",
    "仓库首页": "29-warehouse-home",
    "裁片库": "30-warehouse-cutting-piece",
    "面料库": "31-warehouse-fabric",
    "成品库": "32-warehouse-finished",
    "辅料库": "33-warehouse-accessory",
    "入库": "30-warehouse-cutting-piece",
    "出库": "30-warehouse-cutting-piece",
    "库存": "30-warehouse-cutting-piece",
    "台账": "30-warehouse-cutting-piece",
    "盘点": "30-warehouse-cutting-piece",
    "仓管": "29-warehouse-home",
    "产能配置": "34-capacity-config",
    "产能": "34-capacity-config",
    "工作日历": "35-work-calendar",
    "日历": "35-work-calendar",
    "排产策略": "36-strategy",
    "策略": "36-strategy",
    "系统参数": "37-system-params",
    "操作日志": "38-operation-logs",
    "日志": "38-operation-logs",
    "用户管理": "39-user-management",
    "实际复核": "40-actual-review",
    "实际产量复核": "40-actual-review",
    "actual-review": "40-actual-review",
    "产量复核": "40-actual-review",
    "个人设置": "41-user-settings",
    "设置": "41-user-settings",
    "修改密码": "41-user-settings",
    "整体布局": "01-dashboard",
    "主界面": "01-dashboard",
    "登录": "00-login",
    "login": "00-login",
    "403": "42-403",
    "工作台入口": "43-home-entry",
}

# ASCII 框图字符:以 ┌─┐│└┘├┤┬┴┼ 或 +-+| 开头
ASCII_RE = re.compile(
    r"^```\s*\n([\s\S]*?(?:[┌┐└┘├┤┬┴┼─│]|^[ \t]*\+[-+]+\+|.*\|.*\|.*$|\|.*─.*\|)[\s\S]*?)\n```\s*$",
    re.M,
)


def find_route_key(text_before, current_heading="", ascii_first_line=""):
    """从上下文文本找最匹配的路由 key
    策略:长关键词优先(避免通用词"工作台"先匹配)
    """
    haystack = (text_before + " " + current_heading + " " + ascii_first_line).lower()
    best = None
    best_score = 0
    # 按 keyword 长度降序,长关键词先匹配
    sorted_keys = sorted(ROUTE_PNG.items(), key=lambda x: -len(x[0]))
    for keyword, png_key in sorted_keys:
        if keyword.lower() in haystack:
            score = len(keyword)
            if score > best_score:
                best_score = score
                best = png_key
    return best


def detect_role_for_png(png_key, current_role=None):
    """根据 png_key 和当前角色,选合适截图文件"""
    # 角色相关:工作台/首页在不同角色下指向不同页
    if current_role and current_role.startswith("dispatcher"):
        if png_key in ("01-dashboard", "43-home-entry"):
            png_key = "19-dispatch-home"  # 报工员工作台 = 报工首页
        elif png_key == "42-403":
            pass  # 403 还是 403
    # 角色→子目录(优先当前角色)
    role_dirs = [current_role, "admin", "manager", "planner", "supervisor_sewing",
                 "dispatcher_cutting", "dispatcher_sewing"]
    for role_dir in role_dirs:
        if not role_dir:
            continue
        p = IMG_SCREENSHOT / role_dir / f"{png_key}.png"
        if p.exists():
            return p
    return None


def replace_mermaid_and_ascii(text, current_role=None, current_heading=""):
    """预处理 markdown:
       1. mermaid ```mermaid 块 → ![alt](png)
       2. ASCII 框图 ``` 块 → ![alt](png,智能匹配)
    """
    # 累积"前文"用于 ASCII 框图匹配
    out_parts = []
    cursor = 0
    text_before = ""

    # 1. Mermaid 块替换
    def merm_repl(m):
        nonlocal text_before
        # 占位用文本
        block = m.group(0)
        text_before += block + "\n"
        # 找对应的 PNG(用 mermaid 块前的 heading 序号推断)
        prefix_match = re.search(r"^##\s+", text_before[:text_before.rfind(block)][::-1])
        return f"\n[[MERMAID_PLACEHOLDER:{m.start()}]]\n"

    text = re.sub(r"^```mermaid\s*\n[\s\S]*?\n```\s*$", merm_repl, text, flags=re.M)

    return text


def md_to_docx(md_path, out_path, current_role=None):
    """主转换"""
    import sys
    print(f"  start parse {md_path.name}", flush=True)
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    rpr = style.element.rPr
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rfonts.set(qn('w:ascii'), 'Microsoft YaHei')

    # 页边距
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    text = md_path.read_text(encoding="utf-8")
    print(f"  parsed text {len(text)} chars", flush=True)

    # 切分为块(行级扫描)
    lines = text.split("\n")
    i = 0
    n = len(lines)
    img_count = 0
    tbl_count = 0
    import time
    last_print = time.time()
    while i < n:
        line = lines[i]
        s = line.rstrip()
        if time.time() - last_print > 3:
            print(f"  ... line {i}/{n} s={s[:40]!r}", flush=True)
            last_print = time.time()

        # 空行
        if not s.strip():
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.+)$", s)
        if m:
            level = len(m.group(1))
            content = m.group(2).strip()
            p = doc.add_heading("", level=min(level, 4))
            run = p.add_run(content)
            run.font.name = "Microsoft YaHei"
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            i += 1
            continue

        # Mermaid 块
        if s.strip() == "```mermaid":
            # 找结束
            j = i + 1
            while j < n and lines[j].strip() != "```":
                j += 1
            # 提取块内容
            block_content = "\n".join(lines[i+1:j])
            # 找对应的 PNG(用出现顺序编号)
            out_parts_idx = mermaid_index.get(md_path.name, 0)
            png = IMG_MERMAID / f"{md_to_docx.sop_key(md_path.name)}__{out_parts_idx+1:02d}.png"
            mermaid_index[md_path.name] = out_parts_idx + 1
            if png.exists():
                add_image_paragraph(doc, png, max_width_in=6.5)
                img_count += 1
            else:
                doc.add_paragraph(f"[mermaid PNG 缺失: {png.name}]")
            i = j + 1
            continue

        # 代码块
        if s.strip().startswith("```"):
            lang = s.strip()[3:].strip()
            j = i + 1
            while j < n and not lines[j].strip().startswith("```"):
                j += 1
            code_text = "\n".join(lines[i+1:j])
            # ASCII 框图识别
            is_ascii = any(c in code_text for c in "┌┐└┘├┤┬┴┼─│") or \
                       re.search(r"^\s*[\+\-]+\s*[\+\-]+", code_text, re.M) is not None
            if is_ascii:
                # 智能匹配真实截图
                # 上下文来源:① 最近 ##/### 标题 ② 框图内首行 ③ 框图前 3 行
                first_line = code_text.split("\n", 1)[0].strip() if code_text else ""
                # 框图内 2-3 行(可能是标题)
                inside_lines = [l.strip() for l in code_text.split("\n") if l.strip() and not re.match(r"^[\s│┌┐└┘├┤┬┴┼─]+$", l)]
                inside_text = " ".join(inside_lines[:3])
                nearest_heading = find_nearest_heading(lines, i)
                # 框图前 3 行(避免全文 context 过宽)
                recent_context = extract_context(lines, i, 3)
                route_key = find_route_key(
                    text_before=recent_context,
                    current_heading=nearest_heading,
                    ascii_first_line=first_line + " " + inside_text,
                )
                if False:  # DEBUG
                    import sys as _s
                    _s.stdout.buffer.write(f"      DEBUG heading={nearest_heading!r} inside={inside_text[:50]!r} -> route_key={route_key!r}\n".encode("utf-8", errors="replace"))
                    _s.stdout.buffer.flush()
                if route_key:
                    png = detect_role_for_png(route_key, current_role)
                    if png:
                        # 提取框图上方的"图 N: xxx"作为 caption
                        cap = find_caption(lines, i)
                        if not cap:
                            cap = f"对应系统页面: {route_key}"
                        add_image_paragraph(doc, png, caption=cap, max_width_in=6.5)
                        img_count += 1
                        i = j + 1
                        continue
                # 没匹配到,降级为代码块(等宽字体,加 caption 提示)
                cap = find_caption(lines, i)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[ASCII 示意图{(' - ' + cap) if cap else ''}]")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                i = j + 1
                continue
            else:
                # 普通代码块
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                i = j + 1
                continue

        # 表格
        if "|" in s and i + 1 < n and re.match(r"^\s*\|?[\s\-:|]+\|", lines[i+1]):
            tbl_lines = []
            while i < n and "|" in lines[i]:
                tbl_lines.append(lines[i])
                i += 1
            add_table(doc, tbl_lines)
            tbl_count += 1
            continue

        # 引用
        if s.startswith(">"):
            j = i
            quote_lines = []
            while j < n and lines[j].startswith(">"):
                quote_lines.append(lines[j].lstrip("> ").rstrip())
                j += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            run = p.add_run("\n".join(quote_lines))
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i = j
            continue

        # 列表
        if re.match(r"^(\s*)[-*+]\s+", s) or re.match(r"^(\s*)\d+\.\s+", s):
            j = i
            list_items = []
            while j < n and (re.match(r"^(\s*)[-*+]\s+", lines[j]) or re.match(r"^(\s*)\d+\.\s+", lines[j])):
                list_items.append(lines[j])
                j += 1
            add_list(doc, list_items)
            i = j
            continue

        # 水平线
        if re.match(r"^[\-*_]{3,}$", s.strip()):
            doc.add_paragraph("─" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # 普通段落
        # 累积到段落结束
        j = i
        para_lines = []
        while j < n and lines[j].strip() and not is_block_start(lines[j]):
            para_lines.append(lines[j])
            j += 1
        if para_lines:
            add_paragraph(doc, " ".join(para_lines).strip())
        i = j

    doc.save(str(out_path))
    print(f"  saved {out_path.name} ({img_count} images, {tbl_count} tables)", flush=True)


def is_block_start(line):
    s = line.strip()
    if not s: return True
    if re.match(r"^#{1,6}\s+", s): return True
    if s.startswith("```"): return True
    if "|" in s and re.match(r"^\s*\|?[\s\-:|]+\|", s): return True
    if s.startswith(">"): return True
    if re.match(r"^(\s*)[-*+]\s+", s): return True
    if re.match(r"^(\s*)\d+\.\s+", s): return True
    if re.match(r"^[\-*_]{3,}$", s): return True
    return False


def add_paragraph(doc, text):
    """处理 inline 格式: **bold** *italic* `code` [text](url) ![](url)"""
    p = doc.add_paragraph()
    # 按 token 切分
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|!\[[^\]]*\]\([^)]+\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif part.startswith("![") and "](" in part:
            # 嵌入图片
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", part)
            if m:
                alt, url = m.groups()
                # 解析 url: 相对路径
                if not os.path.isabs(url):
                    url = str(SOP_DIR / url)
                if os.path.exists(url):
                    add_image_paragraph(doc, Path(url), caption=alt, max_width_in=6.0)
                    continue
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                txt, url = m.groups()
                run = p.add_run(txt)
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                run.underline = True
                continue
        else:
            p.add_run(part)


def add_image_paragraph(doc, img_path, caption="", max_width_in=6.5):
    """插入图片(限制宽度),并附 caption"""
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(max_width_in))
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs:
                r.font.size = Pt(9)
                r.italic = True
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    except Exception as e:
        doc.add_paragraph(f"[图片加载失败: {img_path.name}: {e}]")


def add_table(doc, tbl_lines):
    """解析 markdown 表格,转 docx 表格"""
    if len(tbl_lines) < 2:
        return
    # 解析行
    rows = []
    for tl in tbl_lines:
        cells = [c.strip() for c in tl.strip().strip("|").split("|")]
        rows.append(cells)
    # 第二行是 separator
    if len(rows) >= 2 and re.match(r"^[\s\-:|]+$", "|".join(rows[1])):
        header = rows[0]
        data = rows[2:]
    else:
        header = rows[0]
        data = rows[1:]
    if not header:
        return
    ncols = len(header)
    table = doc.add_table(rows=1 + len(data), cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 写表头
    for j, h in enumerate(header):
        if j < ncols:
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.name = "Microsoft YaHei"
    # 写数据
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            if j < ncols:
                cell = table.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(9)
                run.font.name = "Microsoft YaHei"
    doc.add_paragraph("")


def add_list(doc, items):
    """列表"""
    for item in items:
        m = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.+)$", item)
        if not m:
            continue
        indent_spaces, marker, content = m.groups()
        indent = len(indent_spaces) // 2
        try:
            p = doc.add_paragraph(style="List Bullet" if marker in "-*+" else "List Number")
        except KeyError:
            p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6 + indent * 0.4)
        # 简单处理 - 复杂嵌套略
        add_paragraph(doc, content.strip())


def extract_context(lines, i, before_n=10):
    """提取 ASCII 框图前的 N 行文本"""
    start = max(0, i - before_n)
    return "\n".join(lines[start:i])


def find_nearest_heading(lines, i):
    """找 ASCII 框图上方最近的 ## 或 ### 标题(白话内容)"""
    for k in range(i - 1, -1, -1):
        s = lines[k].strip()
        if re.match(r"^#{1,6}\s+", s):
            return s.lstrip("#").strip()
    return ""


def find_caption(lines, i, lookback=5):
    """找 ASCII 框图前 N 行内的「图 N: xxx」说明"""
    for k in range(i - 1, max(0, i - lookback) - 1, -1):
        s = lines[k].strip()
        if not s:
            continue
        m = re.match(r"^图\s*\d+[:：]\s*(.+)$", s)
        if m:
            return m.group(1).strip()
        # 找到非空行但不是图说明,停
        if not s.startswith("图") and not s.startswith("表"):
            break
    return ""


# mermaid 块索引(全局累计)
mermaid_index = {}


def sop_key(fname):
    """SOP 文件名 → 短 key"""
    m = re.match(r"^(SOP-\d{2})", fname)
    return m.group(1) if m else fname.replace(".md", "")


# 给 md_to_docx 函数加 sop_key 方法
md_to_docx.sop_key = staticmethod(sop_key)


SOP_FILES = [
    ("SOP-00-总览与登录.md", None),
    ("SOP-01-系统管理员.md", "admin"),
    ("SOP-02-计划主管.md", "manager"),
    ("SOP-03-计划员.md", "planner"),
    ("SOP-04-报工人员.md", "dispatcher_cutting"),
    ("SOP-05-裁片仓管员.md", "manager"),
    ("SOP-06-缝制车间主任.md", "supervisor_sewing"),
]


if __name__ == "__main__":
    target = sys.argv[1] if len(sys.argv) > 1 else None
    for fname, role in SOP_FILES:
        if target and target not in fname:
            continue
        md_path = SOP_DIR / fname
        if not md_path.exists():
            print(f"SKIP {fname}")
            continue
        out_path = SOP_DIR / f"{fname.replace('.md', '')}.docx"
        print(f"[{fname}] start (role={role})", flush=True)
        mermaid_index[fname] = 0
        try:
            md_to_docx(md_path, out_path, current_role=role)
            size_kb = out_path.stat().st_size // 1024
            print(f"[{fname}] OK ({size_kb}KB)", flush=True)
        except Exception as e:
            print(f"[{fname}] FAIL: {e}", flush=True)
            import traceback
            traceback.print_exc()
