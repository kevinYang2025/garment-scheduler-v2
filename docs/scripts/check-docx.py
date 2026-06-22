"""检查 docx 内每张图片前的 caption 和所在段落(用 style 找 heading)"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import os

for fname in [
    "SOP-00-总览与登录.docx",
    "SOP-04-报工人员.docx",
    "SOP-05-裁片仓管员.docx",
]:
    p = os.path.join("docs/sop", fname)
    if not os.path.exists(p):
        continue
    d = Document(p)
    print(f"\n=== {fname} ===")
    last_heading = ""
    for i, para in enumerate(d.paragraphs):
        text = para.text.strip()
        # 用 style 名找 heading
        if para.style.name.startswith("Heading"):
            last_heading = text[:30]
        # 找图片
        if para._p.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
            cap = ""
            for k in range(i + 1, min(i + 4, len(d.paragraphs))):
                t = d.paragraphs[k].text.strip()
                if t.startswith("对应系统") or t.startswith("图 "):
                    cap = t[:50]
                    break
            print(f"  IMG @ heading='{last_heading}' caption='{cap}'")
